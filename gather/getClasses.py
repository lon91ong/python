# -*- coding: utf-8 -*-
"""
Created on Sat Apr 27 15:22:33 2019
@author: xiaoniu29

getClassName：
从课程表读取年级、班级简称信息，从网络补全班级简称
返回：年级，班级列表
"""
import re, sys
from requests import get as browser
from requests.exceptions import RequestException
from myMod import mbox

def getClassName(docpth):
    from docx import Document
    try:
        word = Document(docpth)
        grade = re.search(r'(?<=20)\d{2}(?=级)',word.paragraphs[0].text).group()
        table = word.tables[0]
        classes = []
        for i in range(1,len(table.rows)):
            for j in range(2,len(table.columns)):
                mach = re.match(r'[\u4e00-\u9fa5]{2,5}\d{1,2}(?=班)',table.cell(i,j).text)
                if mach != None:
                    classes.append(mach.group())
    except:
        #print('提取班级信息失败，请查验Word文档！')
        mbox('错误','提取班级信息失败，请查验Word文档!','error')
        sys.exit()
    #print(classes)
    
    try:
        req = browser('http://211.81.249.110/hmc/hmc.asp',verify=False,timeout=(0.5,3))
    except RequestException as err:
        #print("Error class:",sys.exc_info()[0])
        mbox('错误','网络连接错误:{0}\n错误类型:{1}\n请查验后重试!'.format(err,type(err)),'error')
        sys.exit(0)
    content = req.content.decode('gbk', 'ignore')
    allCla = re.findall(r'[\u4e00-\u9fa5]{2,10}(?=' + grade + '-\d{1,2}</a>)', content)
    allCla = list(set(allCla)) #合并重复，只保留专业全称
    #print(len(allCla))
    
    for ci in classes:
        temp = []
        for j in range(len(allCla)):
            n = len(ci)-1   #尾数不参与比对
            for c in ci[:-1]: #简称所有字符都在全称内
                if (c in allCla[j]):
                    n=n-1
                else:
                    break
            if n==0: #全部找到用全称替换简称
                temp.append(allCla[j])
        if len(temp)==0: # 无匹配时剔除该班级
            mbox('警告','{0}无匹配，请查验！'.format(ci),'warn')
            classes.remove(ci)
        else:
            classes[classes.index(ci)] = (min(temp,key=len)+grade+'-'+ci[-1]) # 存在多个匹配对象时，取最短的
    return grade, classes
