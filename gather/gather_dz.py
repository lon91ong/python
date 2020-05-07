# -*- coding: utf-8 -*-
"""
Created on Fri Apr  5 14:50:32 2019
抓取教务网站学生花名册, 三楼适用(多年级,单纯电子表格输出)
@author: xiaoniu29
"""

import sys, re
from os import path
from time import localtime
from xlwt import Workbook
from requests import get as browser
from requests.exceptions import RequestException
from urllib.parse import quote

from myMod import mbox, GetDesktopPath, getFile
from getClasses import getClassName

def getClassName(docpth):
    from docx import Document
    
    try:
        req = browser('http://211.81.249.110/hmc/hmc.asp',verify=False,timeout=(0.5,3))
    except RequestException as err:
        #print("Error class:",sys.exc_info()[0])
        mbox('错误','网络连接错误:{0}\n错误类型:{1}\n请查验后重试!'.format(err,type(err)),'error')
        sys.exit(0)
    content = req.content.decode('gbk', 'ignore')
    
    classes = []
    try:
        word = Document(docpth)
        #print('LEN:',len(word.paragraphs))
        for pg in word.paragraphs:
            #print(pg.text)
            grade = re.search(r'\d{2}(?=级)',pg.text)
            if grade:
                grade = grade.group()
            else:
                continue
            pattern = re.compile(r'(?<=\d{2}级)[一-龥]{2,9}\d{1,2}(?=班)',re.S|re.U)
            mach = pattern.findall(pg.text)
            #mach = re.finditer(r'(?<=\d{2}级)[一-龥]{2,9}\d{1,2}(?=班)',pg.text)
            #print('年级:',grade,'LEN:',len(mach))
            if len(mach) > 0:
                for mc in mach:
                    mc = mc[:-1]+grade+'-'+mc[-1]
                    classes.append(mc)
                #print(mach)
                #classes += mach
    except:
        print("Unexpected error:", sys.exc_info())
        mbox('错误','提取班级信息失败，请查验Word文档(DocX Only)!','error')
        sys.exit()
    #print(classes)

    allCla = re.findall(r'[一-龥]{2,10}(?=\d{2}-\d{1,2}</a>)', content)
    allCla = list(set(allCla)) #合并重复，只保留专业全称
    #print(len(allCla))
    
    for ci in classes:
        temp = []
        for j in range(len(allCla)):
            n = len(ci)-4   #尾数不参与比对
            for c in ci[:-4]: #简称所有字符都在全称内
                if (c in allCla[j]):
                    n=n-1
                else:
                    break
            if n==0: #全部找到用全称替换简称
                temp.append(allCla[j])
        if len(temp)==0: # 无匹配时剔除该班级
            mbox('警告','{0}无匹配，请查验！'.format(ci),'warn')
            classes.remove(ci)
        else:
            classes[classes.index(ci)] = (min(temp,key=len)+ci[-4:]) # 存在多个匹配对象时，取最短的
    return classes

def main(argv):
    exedir = path.dirname(path.abspath(sys.argv[0])) # 程序所在目录
    if len(sys.argv) < 2: #直接运行程序
        docpth = getFile(exedir,'打开课程表(仅支持docx格式)',"Word 10文档 (*.docx)|*.docx||")
        if docpth == '':
            mbox( '错误','需要指定课程表文件!', 'error')
            sys.exit(0)
    else:
        docpth = path.abspath(sys.argv[-1])
    classes = getClassName(docpth) #取班级列表
    #print(classes)

    #花名电子表
    wb = Workbook(encoding="utf-8")
    worksheet = wb.add_sheet("花名", cell_overwrite_ok=True)
    worksheet.write(0, 0,'班级')
    students = []
    n = 0
    for i in range(len(classes)):
        worksheet.write(i+1,0,classes[i])
        try:
            req = browser('http://211.81.249.110/hmc/hmc_p.asp?trbj='+quote(classes[i],encoding='gb2312'),verify=False,timeout=(0.5,3))
        except RequestException as err:
            #print("Error class:",sys.exc_info()[0])
            mbox('错误','网络连接错误:{0}\n错误类型:{1}\n请查验后重试!'.format(err,type(err)),'error')
            sys.exit(0)
        #content = req.content.decode('gbk', 'ignore') # 忽略非法字符，replace则用?取代非法字符；
        #res = re.findall(r'\d{12}(?=</td>)', content) # 所有的学号
        res = re.findall(r'(\d{12})</td>\s+<td align="left">&nbsp;([一-龥]{2,5})</td>', req.content.decode('gbk', 'ignore'), re.S) # 成对提取学号和姓名
        if len(res) == 0:
            mbox('获取名单失败!','无法获取{}班名单，请手动检查教务名单页面！'.format(classes[i]),'error')
            sys.exit(0)
        #sum(['序号{0},成绩{0}'.format(i).split(',') for i in range(1, len(res)+1)],['班级'])
        for j in range(1,len(res)+1):
            students += [(res[j-1][0],res[j-1][1],classes[i])]
            worksheet.write(i+1,j,res[j-1][1])
            if j > n: worksheet.write(0,j,'序号'+str(j))
        if n < j: n = j
        #students += list((x[0],x[1],classes[i]) for x in res)

    wb.save(GetDesktopPath()+'\\muster.xls') #桌面路径
    mbox('完成','请在桌面查看muster.xls文件!','info')

if __name__ == "__main__":
    main(sys.argv[1:])
