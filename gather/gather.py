# -*- coding: utf-8 -*-
"""
Created on Fri Apr  5 14:50:32 2019
抓取教务网站学生花名册
@author: xiaoniu29
"""

import sys, re, requests
from lxml import etree 
import sqlite3, os
import docx, xlwt
from urllib.parse import quote

def mbox(title, text, style = ''):
    import win32api,win32con
    if style == 'error':
        win32api.MessageBox(0, text, title, win32con.MB_ICONERROR)
    elif style == 'info':
        win32api.MessageBox(0, text, title, win32con.MB_ICONASTERISK)
    elif style == 'warn':
        win32api.MessageBox(0, text, title, win32con.MB_ICONWARNING)
    else:
        win32api.MessageBox(0, text, title, win32con.MB_OK)

def GetDesktopPath():
    import winreg
    key = winreg.OpenKey(winreg.HKEY_CURRENT_USER,r'Software\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders')
    return winreg.QueryValueEx(key, "Desktop")[0]

def main(argv):
    try:
        word = docx.Document(str(argv[-1]))
        #print('接收到的文件：'+str(argv[-1]))
        #word = docx.Document(r'E:\花名\M2019年春季学期2017级第二轮.docx')
        grade = re.search(r'(?<=20)\d{2}(?=级)',word.paragraphs[0].text).group()
        table = word.tables[0]
        classes = []
        for i in range(1,len(table.rows)):
            for j in range(2,len(table.columns)):
                mach = re.match(r'[\u4e00-\u9fa5]{2,5}\d{1,2}(?=班)',table.cell(i,j).text)
                if mach != None:
                    classes.append(mach.group())
    except:
        #print('提取班级信息失败，请查验Word文档！')
        mbox('错误','提取班级信息失败，请查验Word文档!','error')
        sys.exit()
    #print(classes)
    
    try:
        req = requests.get('http://211.81.249.110/hmc/hmc.asp')
    except ConnectionResetError as err:
        mbox('错误','网络连接错误:{0}\n请查验Word文档!'.format(err),'error')
        sys.exit(0)
    content = req.content.decode('gbk', 'ignore')
    allCla = re.findall(r'[\u4e00-\u9fa5]{2,10}(?=' + grade + '-\d{1,2}</a>)', content)
    allCla = list(set(allCla)) #合并重复，只保留专业全称
    #print(len(allCla))
    
    for i in range(len(classes)):
        temp = []
        for j in range(len(allCla)):
            n = len(classes[i])-1   #尾数不参与比对
            for c in classes[i][:-1]: #简称所有字符都在全称内
                if not(c in allCla[j]):
                    break
                else:
                    n = n-1
            if n==0: #全部找到用全称替换简称
                temp.append(allCla[j])
        #if len(temp)>1:print(classes[i],temp)
        classes[i] = min(temp,key=len)+grade+'-'+classes[i][-1] # 存在多个匹配对象时，取最短的
    #print(classes)
    
    
    conn = sqlite3.connect(os.path.dirname(os.path.abspath(sys.argv[0]))+'/hwmy'+grade+'.db') #参数路径
    #mbox('info',os.path.dirname(os.path.realpath(sys.argv[0])),'info')
    #conn = sqlite3.connect(':memory:') #内存临时
    curs = conn.cursor()
    conn.text_factory = str #lambda x: str(x, "gbk", "ignore")
    try:
        curs.execute('select * from sqlite_master where type = "table" and name = "students"')
        if curs.fetchone() == None:
            curs.execute('create table students (id varchar(12) primary key, name varchar(12) not NULL collate nocase, class text collate nocase, score unsigned tinyint, unique (id))')
        curs.execute('select * from sqlite_master where type = "table" and name = "classes"')
        if curs.fetchone() == None:
            curs.execute('create table classes (id INTEGER PRIMARY KEY AUTOINCREMENT, class text not NULL collate nocase, unique (class))')
    except sqlite3.OperationalError as e:
        print("Error info:",e.args[0])
        pass
    #curs.executemany('INSERT INTO classes VALUES (?,?,?)',[(3,'name3',19),(4,'name4',26)])
    
    #花名电子表
    workbook=xlwt.Workbook(encoding="utf-8")
    worksheet = workbook.add_sheet("花名", cell_overwrite_ok=True)
    worksheet.write(0, 0,'班级')
    for i in range(len(classes)):
        worksheet.write(i+1,0,classes[i])
        
        try:
            curs.execute('insert into classes (class) values ("{0}")'.format(classes[i]))
        except (sqlite3.IntegrityError, sqlite3.OperationalError) as e:
            print("Error info:",e.args[0])
            pass
        
        req = requests.get('http://211.81.249.110/hmc/hmc_p.asp?trbj='+quote(classes[i],encoding='gb2312'))
        content = req.content.decode('gbk', 'ignore') # 忽略非法字符，replace则用?取代非法字符；
        ids = re.findall(r'\d{12}(?=</td>)', content)   # 所有的学号
        root = etree.HTML(content)
        name = ''
        for j in range(len(ids)):
            name = root.xpath('//td[text()='+ids[j]+']/following-sibling::td[1]/text()')[0].lstrip('\xa0')
            worksheet.write(i+1,j+1,name)
            worksheet.write(0,j+1,'序号'+str(j+1))
            
            try:
                curs.execute('INSERT INTO students (id, name, class) VALUES ("%s", "%s", "%s")' % (ids[j],name,classes[i]))
            except (sqlite3.IntegrityError, sqlite3.OperationalError):
                pass
        
    conn.commit()
    
    workbook.save(GetDesktopPath()+'\\muster.xls') #桌面路径
    
    #curs.execute('SELECT name FROM students WHERE id="171304011039"')
    #curs.execute('SELECT * FROM classes)
    #print(curs.fetchone()[0])
    #print(len(curs.fetchall()))
    #curs.execute("drop table classes") #删除表
    curs.close()
    conn.close()
    
    mbox('完成','请在桌面查看muster.xls文件!','info')

if __name__ == "__main__":
   main(sys.argv[1:])
